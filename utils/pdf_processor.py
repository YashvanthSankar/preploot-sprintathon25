import os
import glob
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain_chroma import Chroma
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from config import Config
from utils.user_manager import (
    create_user_directories, 
    get_user_processed_files, 
    save_user_processed_files,
    list_user_files
)
import json

class PDFProcessor:
    def __init__(self, user_id):
        self.user_id = user_id
        self.embeddings = GoogleGenerativeAIEmbeddings(
            model="gemini-embedding-001",
            google_api_key=Config.GOOGLE_API_KEY
        )
        
        # Create user directories
        self.data_folder, self.vectordb_folder = create_user_directories(user_id)
        
    def get_text_from_pdf(self, file_path):
        """Extract text from PDF file."""
        reader = PdfReader(file_path)
        text = ""
        for page in reader.pages:
            content = page.extract_text()
            if content:
                text += content + "\n"
        return text
    
    def get_text_from_docx(self, file_path):
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def get_text_from_file(self, file_path):
        """Extract text from either PDF or DOCX file."""
        if file_path.endswith('.pdf'):
            return self.get_text_from_pdf(file_path)
        elif file_path.endswith('.docx'):
            return self.get_text_from_docx(file_path)
        else:
            return ""
    
    def check_new_files(self):
        """Check for new or updated files for this user."""
        # Get both PDF and DOCX files
        pdf_files = glob.glob(f"{self.data_folder}/*.pdf")
        docx_files = glob.glob(f"{self.data_folder}/*.docx")
        all_files = pdf_files + docx_files
        
        processed_files = get_user_processed_files(self.user_id)

        # Detect removed files
        removed_files = [f for f in list(processed_files.keys()) if f not in all_files]
        for f in removed_files:
            print(f"File removed: {f}")
            processed_files.pop(f, None)

        # Detect new/updated files
        new_files = []
        for file in all_files:
            last_modified = os.path.getmtime(file)
            if (file not in processed_files) or (processed_files[file] < last_modified):
                new_files.append(file)
                processed_files[file] = last_modified

        # Save state immediately after updating processed_files
        save_user_processed_files(self.user_id, processed_files)

        return new_files, removed_files
    
    def store_chroma_function(self):
        """Process PDFs and store in user-specific vector database."""
        new_files, removed_files = self.check_new_files()
        
        # Load existing DB if it exists, otherwise create later
        vectorstore = None
        if os.path.exists(self.vectordb_folder):
            vectorstore = Chroma(persist_directory=self.vectordb_folder, embedding_function=self.embeddings)

        # Delete vectors for removed files
        if removed_files and vectorstore:
            for f in removed_files:
                print(f"Deleting vectors for: {f}")
                vectorstore.delete(where={"source": f})
        
        # Check if there are any PDF or DOCX files at all
        pdf_files = glob.glob(f"{self.data_folder}/*.pdf")
        docx_files = glob.glob(f"{self.data_folder}/*.docx")
        all_files = pdf_files + docx_files
        
        if not all_files:
            print("No PDF or DOCX files in the user directory.")
            if vectorstore:
                print("Cleaning up vector store...")
                vectorstore._collection.delete(where={})
                vectorstore._client._conn.close()
                vectorstore = None
                
            # Delete the vector store directory and its contents
            if os.path.exists(self.vectordb_folder):
                print(f"Removing {self.vectordb_folder} directory...")
                import shutil
                shutil.rmtree(self.vectordb_folder)
                
            state_file = Config.get_user_state_file(self.user_id)
            if os.path.exists(state_file):
                print("Removing state file...")
                os.remove(state_file)
            return None

        # No new files → just return
        if not new_files:
            print("No new files, skipping embeddings.")
            return vectorstore

        # Otherwise, process new files
        print("New/updated files found:", new_files)
        docs = []
        for f in new_files:
            text = self.get_text_from_file(f)
            docs.append(Document(page_content=text, metadata={"source": f}))
        
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=100)
        chunks = text_splitter.split_documents(docs)

        if vectorstore:
            vectorstore.add_documents(chunks)
        else:
            vectorstore = Chroma.from_documents(documents=chunks, embedding=self.embeddings, persist_directory=self.vectordb_folder)

        processed_files = get_user_processed_files(self.user_id)
        save_user_processed_files(self.user_id, processed_files)
        return vectorstore
    
    def get_vectorstore(self):
        """Get existing vector store for this user."""
        if os.path.exists(self.vectordb_folder):
            return Chroma(
                persist_directory=self.vectordb_folder,
                embedding_function=self.embeddings
            )
        return None
    
    def save_file(self, file_data, filename):
        """Save uploaded file (PDF or DOCX) to user's data folder."""
        file_path = os.path.join(self.data_folder, filename)
        with open(file_path, 'wb') as f:
            f.write(file_data)
        return file_path